from __future__ import annotations

import html
import re
import sys
from pathlib import Path


ROOT = Path(__file__).resolve().parents[2]
DEPS = ROOT / ".codex" / "docdeps"
if DEPS.exists():
    sys.path.insert(0, str(DEPS))

from docx import Document  # type: ignore
from docx.enum.section import WD_SECTION_START  # type: ignore
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH  # type: ignore
from docx.shared import Cm, Pt  # type: ignore
from reportlab.lib import colors  # type: ignore
from reportlab.lib.enums import TA_CENTER, TA_LEFT  # type: ignore
from reportlab.lib.pagesizes import A4  # type: ignore
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore
from reportlab.lib.units import cm  # type: ignore
from reportlab.pdfbase import pdfmetrics  # type: ignore
from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
from reportlab.platypus import (  # type: ignore
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


SOURCE = ROOT / "docs" / "08-legal" / "hanuja-satici-pazaryeri-sozlesmesi.md"
OUT_DIR = ROOT / "output" / "doc"
DOCX_OUT = OUT_DIR / "hanuja-satici-pazaryeri-sozlesmesi.docx"
PDF_OUT = OUT_DIR / "hanuja-satici-pazaryeri-sozlesmesi.pdf"


def read_lines() -> list[str]:
    return SOURCE.read_text(encoding="utf-8").splitlines()


def clean_inline(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def add_docx_inline(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw = lines[idx].strip().strip("|")
        cells = [clean_inline(cell.strip()) for cell in raw.split("|")]
        if not all(set(cell) <= {"-", ":", " "} for cell in cells):
            rows.append(cells)
        idx += 1
    return rows, idx


def build_docx() -> None:
    lines = read_lines()
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True

    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()
        if not stripped or stripped == "---":
            idx += 1
            continue
        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            if rows:
                table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell in enumerate(row):
                        target = table.rows[r_idx].cells[c_idx]
                        target.text = cell
                        for paragraph in target.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = "Arial"
                                run.font.size = Pt(9)
                                if r_idx == 0:
                                    run.bold = True
                document.add_paragraph()
            continue
        if stripped.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_docx_inline(paragraph, stripped[2:])
        elif stripped.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 1")
            add_docx_inline(paragraph, stripped[3:])
        elif stripped.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 2")
            add_docx_inline(paragraph, stripped[4:])
        elif stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_docx_inline(paragraph, stripped[2:])
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.08
            add_docx_inline(paragraph, stripped)
        idx += 1

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_OUT)


def register_fonts() -> tuple[str, str]:
    font_dir = Path("C:/Windows/Fonts")
    regular = font_dir / "arial.ttf"
    bold = font_dir / "arialbd.ttf"
    if regular.exists() and bold.exists():
        pdfmetrics.registerFont(TTFont("AgreementArial", str(regular)))
        pdfmetrics.registerFont(TTFont("AgreementArial-Bold", str(bold)))
        return "AgreementArial", "AgreementArial-Bold"
    return "Helvetica", "Helvetica-Bold"


def pdf_inline(text: str) -> str:
    escaped = html.escape(text)
    escaped = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(r"`(.*?)`", r"<font name='Courier'>\1</font>", escaped)
    return escaped


def build_pdf() -> None:
    regular_font, bold_font = register_fonts()
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "AgreementNormal",
        parent=styles["Normal"],
        fontName=regular_font,
        fontSize=8.7,
        leading=11,
        alignment=TA_LEFT,
        spaceAfter=4,
    )
    title = ParagraphStyle(
        "AgreementTitle",
        parent=normal,
        fontName=bold_font,
        fontSize=15,
        leading=19,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    h1 = ParagraphStyle("AgreementH1", parent=normal, fontName=bold_font, fontSize=11.5, leading=14, spaceBefore=8, spaceAfter=5)
    h2 = ParagraphStyle("AgreementH2", parent=normal, fontName=bold_font, fontSize=10, leading=12, spaceBefore=5, spaceAfter=3)
    small = ParagraphStyle("AgreementSmall", parent=normal, fontSize=8, leading=10)

    lines = read_lines()
    story = []
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()
        if not stripped:
            idx += 1
            continue
        if stripped == "---":
            story.append(Spacer(1, 6))
            idx += 1
            continue
        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            if rows:
                data = [[Paragraph(pdf_inline(cell), small) for cell in row] for row in rows]
                table = Table(data, repeatRows=1, hAlign="LEFT")
                table.setStyle(
                    TableStyle(
                        [
                            ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f3f5")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 5),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                            ("TOPPADDING", (0, 0), (-1, -1), 5),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                        ]
                    )
                )
                story.append(table)
                story.append(Spacer(1, 6))
            continue
        if stripped.startswith("# "):
            if story:
                story.append(PageBreak())
            story.append(Paragraph(pdf_inline(stripped[2:]), title))
        elif stripped.startswith("## "):
            story.append(Paragraph(pdf_inline(stripped[3:]), h1))
        elif stripped.startswith("### "):
            story.append(Paragraph(pdf_inline(stripped[4:]), h2))
        elif stripped.startswith("- "):
            items = []
            while idx < len(lines) and lines[idx].strip().startswith("- "):
                items.append(ListItem(Paragraph(pdf_inline(lines[idx].strip()[2:]), normal), leftIndent=10))
                idx += 1
            story.append(ListFlowable(items, bulletType="bullet", leftIndent=14))
            continue
        else:
            story.append(Paragraph(pdf_inline(stripped), normal))
        idx += 1

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        rightMargin=1.5 * cm,
        leftMargin=1.5 * cm,
        topMargin=1.35 * cm,
        bottomMargin=1.35 * cm,
        title="Hanuja Satici Pazaryeri Katilim ve Hizmet Sozlesmesi",
        author="Hanuja",
    )
    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)


def add_page_number(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 7)
    canvas.drawRightString(A4[0] - 1.5 * cm, 0.8 * cm, f"Sayfa {doc.page}")
    canvas.restoreState()


def main() -> None:
    build_docx()
    build_pdf()
    print(DOCX_OUT)
    print(PDF_OUT)


if __name__ == "__main__":
    main()
